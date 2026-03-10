import asyncio
import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from fastapi import HTTPException

from server.agents import references


def test_search_tfidf_fallback(tmp_path, monkeypatch):
    root = tmp_path / "references"
    root.mkdir()
    d = root / "testref"
    d.mkdir()
    pages = [
        {"page": 1, "text": "The fireball spell deals 8d6 fire damage in an area.", "snippet": "fireball deals 8d6 fire damage"},
        {"page": 2, "text": "Short rest: spend hit dice to recover HP.", "snippet": "Short rest rules"},
    ]
    (d / "pages.json").write_text(json.dumps(pages), encoding="utf-8")
    meta = {"title": "Test Reference", "filename": "testref.pdf", "pages": len(pages)}
    (d / "metadata.json").write_text(json.dumps(meta), encoding="utf-8")

    # Point the module storage root at our temp dir
    monkeypatch.setattr(references, "_storage_root", lambda: root)

    results = references.search_query("fireball", top_k=3)
    assert isinstance(results, list)
    assert len(results) >= 1
    assert results[0]["source_id"] == "testref"


def test_list_references(tmp_path, monkeypatch):
    root = tmp_path / "references"
    root.mkdir()
    d = root / "refA"
    d.mkdir()
    meta = {"title": "Ref A", "filename": "a.pdf", "pages": 1}
    (d / "metadata.json").write_text(json.dumps(meta), encoding="utf-8")

    monkeypatch.setattr(references, "_storage_root", lambda: root)

    # list_references is async; run it
    current_user = SimpleNamespace(email='tester@example.com', username='tester', profile=None, role='player')
    out = asyncio.run(references.list_references(current_user=current_user))
    assert isinstance(out, list)
    assert any(item.get("id") == "refA" for item in out)


# ---------------------------------------------------------------------------
# Tests for multi-format text extraction
# ---------------------------------------------------------------------------

def test_extract_txt(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Tavern rules:\n1. No fighting.\n2. Pay your tab.", encoding="utf-8")
    chunks = references._extract_pages_text(f)
    assert len(chunks) >= 1
    assert "Tavern rules" in chunks[0]


def test_extract_markdown(tmp_path):
    f = tmp_path / "readme.md"
    f.write_text("# Campaign Notes\n\nThe dark forest is dangerous.", encoding="utf-8")
    chunks = references._extract_pages_text(f)
    assert len(chunks) >= 1
    assert "Campaign Notes" in chunks[0]


def test_extract_html(tmp_path):
    f = tmp_path / "lore.html"
    f.write_text(
        "<html><body><h1>World Lore</h1><p>The realm was created eons ago.</p>"
        "<script>alert('ignored')</script></body></html>",
        encoding="utf-8",
    )
    chunks = references._extract_pages_text(f)
    assert len(chunks) >= 1
    text = chunks[0]
    assert "World Lore" in text
    assert "realm" in text
    assert "alert" not in text  # script content must be stripped


def test_extract_json(tmp_path):
    data = {"name": "Dragon", "hp": 300, "abilities": ["fire breath", "fly"]}
    f = tmp_path / "monster.json"
    f.write_text(json.dumps(data), encoding="utf-8")
    chunks = references._extract_pages_text(f)
    assert len(chunks) >= 1
    assert "Dragon" in chunks[0]


def test_extract_csv(tmp_path):
    f = tmp_path / "items.csv"
    f.write_text("Name,Type,Cost\nSword,Weapon,15gp\nShield,Armor,10gp", encoding="utf-8")
    chunks = references._extract_pages_text(f)
    assert len(chunks) >= 1
    assert "Sword" in chunks[0]


def test_extract_docx(tmp_path):
    import docx  # python-docx

    doc_path = tmp_path / "adventure.docx"
    doc = docx.Document()
    doc.add_paragraph("Chapter 1: The Beginning")
    doc.add_paragraph("The hero set out on a quest for glory.")
    doc.save(str(doc_path))

    chunks = references._extract_pages_text(doc_path)
    assert len(chunks) >= 1
    assert "hero" in chunks[0]


def test_extract_xlsx(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Treasure"
    ws.append(["Item", "Value"])
    ws.append(["Gold Ring", "50gp"])
    wb.save(str(tmp_path / "treasure.xlsx"))

    chunks = references._extract_pages_text(tmp_path / "treasure.xlsx")
    assert len(chunks) >= 1
    assert "Gold Ring" in chunks[0]


def test_supported_extensions_set():
    """Ensure all expected formats are present in the supported set."""
    expected = {".pdf", ".txt", ".md", ".csv", ".json", ".html", ".htm", ".docx", ".doc", ".xlsx", ".xls"}
    assert expected.issubset(references.SUPPORTED_EXTENSIONS), (
        f"Missing extensions: {expected - references.SUPPORTED_EXTENSIONS}"
    )


# ---------------------------------------------------------------------------
# Folder management
# ---------------------------------------------------------------------------

def test_validate_ref_folder_path_valid():
    """_validate_ref_folder_path should accept well-formed paths and strip leading/trailing slashes."""
    assert references._validate_ref_folder_path("rules") == "rules"
    assert references._validate_ref_folder_path("/rules/combat/") == "rules/combat"
    assert references._validate_ref_folder_path("rules/combat") == "rules/combat"
    assert references._validate_ref_folder_path("") == ""


def test_validate_ref_folder_path_rejects_traversal():
    """_validate_ref_folder_path must block path traversal and invalid segments."""
    for bad in ("..", "rules/../etc", "rules/./combat", "a\\b", "a\0b", "a:b"):
        with pytest.raises(HTTPException) as exc_info:
            references._validate_ref_folder_path(bad)
        assert exc_info.value.status_code == 400


def test_create_and_list_ref_folders(tmp_path, monkeypatch):
    """create_ref_folder persists a folder; list_ref_folders returns it."""
    root = tmp_path / "refs"
    root.mkdir()
    monkeypatch.setattr(references, "_storage_root", lambda: root)

    admin_user = SimpleNamespace(email="admin@example.com", username="admin", profile=None, role="admin")
    monkeypatch.setattr(references._db, "is_admin_user", lambda u: True)

    body = references.FolderItem(folder="combat")
    result = asyncio.run(references.create_ref_folder(body=body, current_user=admin_user))
    assert result["folder"] == "combat"

    out = asyncio.run(references.list_ref_folders(current_user=admin_user))
    assert "combat" in out["folders"]


def test_delete_ref_folder_blocked_by_ref(tmp_path, monkeypatch):
    """delete_ref_folder should return 400 when a reference is assigned to the folder."""
    root = tmp_path / "refs"
    root.mkdir()
    ref_dir = root / "myref"
    ref_dir.mkdir()
    (ref_dir / "metadata.json").write_text(json.dumps({"title": "T", "filename": "f.pdf", "pages": 1, "folder": "combat"}), encoding="utf-8")

    monkeypatch.setattr(references, "_storage_root", lambda: root)
    monkeypatch.setattr(references._db, "is_admin_user", lambda u: True)

    admin_user = SimpleNamespace(email="admin@example.com", username="admin", profile=None, role="admin")

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(references.delete_ref_folder(folder_path="combat", current_user=admin_user))
    assert exc_info.value.status_code == 400


def test_delete_ref_folder_blocked_by_subfolder_ref(tmp_path, monkeypatch):
    """delete_ref_folder should return 400 when a reference exists in a subfolder of the target."""
    root = tmp_path / "refs"
    root.mkdir()
    ref_dir = root / "deepref"
    ref_dir.mkdir()
    (ref_dir / "metadata.json").write_text(json.dumps({"title": "T", "filename": "f.pdf", "pages": 1, "folder": "rules/combat"}), encoding="utf-8")

    monkeypatch.setattr(references, "_storage_root", lambda: root)
    monkeypatch.setattr(references._db, "is_admin_user", lambda u: True)

    admin_user = SimpleNamespace(email="admin@example.com", username="admin", profile=None, role="admin")

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(references.delete_ref_folder(folder_path="rules", current_user=admin_user))
    assert exc_info.value.status_code == 400


def test_move_reference_updates_folder(tmp_path, monkeypatch):
    """move_reference should update the folder field in metadata.json."""
    root = tmp_path / "refs"
    root.mkdir()
    ref_dir = root / "myref"
    ref_dir.mkdir()
    (ref_dir / "metadata.json").write_text(json.dumps({"title": "Ref", "filename": "ref.pdf", "pages": 2, "folder": "old"}), encoding="utf-8")

    monkeypatch.setattr(references, "_storage_root", lambda: root)
    monkeypatch.setattr(references._db, "is_admin_user", lambda u: True)

    admin_user = SimpleNamespace(email="admin@example.com", username="admin", profile=None, role="admin")
    body = references.MoveRefRequest(folder="new-folder")
    result = asyncio.run(references.move_reference(ref_id="myref", body=body, current_user=admin_user))
    assert result["meta"]["folder"] == "new-folder"
    saved = json.loads((ref_dir / "metadata.json").read_text())
    assert saved["folder"] == "new-folder"


def test_move_reference_missing_metadata_uses_defaults(tmp_path, monkeypatch):
    """move_reference should fall back to safe defaults when metadata.json is missing."""
    root = tmp_path / "refs"
    root.mkdir()
    ref_dir = root / "noMeta"  # noqa: F841
    ref_dir.mkdir()
    # No metadata.json written

    monkeypatch.setattr(references, "_storage_root", lambda: root)
    monkeypatch.setattr(references._db, "is_admin_user", lambda u: True)

    admin_user = SimpleNamespace(email="admin@example.com", username="admin", profile=None, role="admin")
    body = references.MoveRefRequest(folder="anywhere")
    result = asyncio.run(references.move_reference(ref_id="noMeta", body=body, current_user=admin_user))
    assert result["meta"]["folder"] == "anywhere"
    assert result["meta"]["title"] == "noMeta"
    assert result["meta"]["filename"] == ""
    assert result["meta"]["pages"] == 0
